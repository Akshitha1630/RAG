"""
Script to generate Word documentation for the RAG PDF Chatbot project.
Run: python generate_docs.py
Output: docs/RAG_PDF_Chatbot_Documentation.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def add_heading_styled(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return heading


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)
    doc.add_paragraph()


def build_document():
    doc = Document()

    # -- Title Page --
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RAG PDF Chatbot")
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Retrieval-Augmented Generation System\nTechnical Documentation")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x77)

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = author.add_run("\n\nRAG PDF Chatbot Project")
    run.font.size = Pt(12)
    doc.add_page_break()

    # -- Table of Contents placeholder --
    add_heading_styled(doc, "Table of Contents", 1)
    toc_items = [
        "1. Introduction", "2. System Architecture", "3. Technology Stack",
        "4. Document Ingestion Pipeline", "5. RAG Query Pipeline",
        "6. Conversation Memory", "7. Project Structure",
        "8. Installation & Setup", "9. Configuration", "10. Usage Guide",
        "11. Troubleshooting", "12. Future Improvements",
    ]
    for item in toc_items:
        doc.add_paragraph(item, style="List Number")
    doc.add_page_break()

    # -- 1. Introduction --
    add_heading_styled(doc, "1. Introduction", 1)
    doc.add_paragraph(
        "This project implements a Retrieval-Augmented Generation (RAG) chatbot "
        "that allows users to upload PDF documents and ask natural language questions "
        "about their content. The system uses a locally-hosted Large Language Model "
        "(LLM) via Ollama, ensuring complete data privacy — no data leaves the user's machine."
    )
    doc.add_paragraph(
        "Unlike traditional chatbots that rely solely on pre-trained knowledge, "
        "RAG combines the power of information retrieval with generative AI. "
        "The system first retrieves the most relevant sections from uploaded documents, "
        "then uses those sections as context for the LLM to generate accurate, "
        "grounded responses."
    )
    add_heading_styled(doc, "Key Features", 2)
    features = [
        "Upload and process PDF documents in real-time through a web interface",
        "Semantic search over document content using vector embeddings",
        "Context-aware responses grounded in actual document content",
        "Multi-turn conversation with memory for follow-up questions",
        "Streaming token-by-token response output for real-time feedback",
        "Persistent vector store — documents survive across app restarts",
        "100% local execution — no API keys or cloud services required",
    ]
    for f in features:
        doc.add_paragraph(f, style="List Bullet")

    # -- 2. System Architecture --
    add_heading_styled(doc, "2. System Architecture", 1)
    doc.add_paragraph(
        "The system follows a two-phase architecture: an ingestion phase that "
        "processes PDF documents into searchable vector embeddings, and a query "
        "phase that retrieves relevant context and generates responses."
    )
    add_heading_styled(doc, "High-Level Data Flow", 2)
    steps = [
        ("PDF Upload", "User uploads a PDF via the Streamlit web interface."),
        ("Text Extraction", "PyPDFLoader extracts raw text from each page of the PDF."),
        ("Text Chunking", "RecursiveCharacterTextSplitter breaks text into 1500-character chunks with 200-character overlap."),
        ("Embedding", "Each chunk is converted to a vector embedding using the Llama2 model via Ollama."),
        ("Storage", "Embeddings are stored in ChromaDB, a persistent vector database on disk."),
        ("User Query", "User types a question in the chat interface."),
        ("Query Embedding", "The question is embedded into the same vector space."),
        ("Similarity Search", "ChromaDB finds the top-4 most similar document chunks."),
        ("Prompt Assembly", "Retrieved chunks + conversation history + question are combined into a prompt."),
        ("LLM Generation", "Llama2 generates a contextually grounded response, streamed to the UI."),
    ]
    add_table(doc, ["Step", "Component", "Description"],
              [(str(i+1), s, d) for i, (s, d) in enumerate(steps)])

    # -- 3. Technology Stack --
    add_heading_styled(doc, "3. Technology Stack", 1)
    stack = [
        ("Streamlit", ">=1.35.0", "Web UI framework for upload and chat interface"),
        ("LangChain", ">=1.3.0", "Orchestration framework for the RAG pipeline"),
        ("LangChain-Ollama", ">=1.1.0", "LLM and embedding integration with Ollama"),
        ("LangChain-Chroma", ">=1.1.0", "Vector store integration with ChromaDB"),
        ("ChromaDB", ">=1.5.0", "Persistent vector database for embeddings"),
        ("Ollama", "—", "Local LLM server hosting the Llama2 model"),
        ("PyPDF", ">=6.0.0", "PDF text extraction library"),
        ("Sentence-Transformers", ">=5.0.0", "Transformer-based embedding models"),
    ]
    add_table(doc, ["Technology", "Version", "Purpose"],
              [(t, v, p) for t, v, p in stack])

    # -- 4. Document Ingestion Pipeline --
    add_heading_styled(doc, "4. Document Ingestion Pipeline", 1)
    add_heading_styled(doc, "4.1 PDF Loading", 2)
    doc.add_paragraph(
        "The PyPDFLoader reads the uploaded PDF file and extracts text content "
        "from each page. It handles multi-page documents and preserves page-level "
        "metadata (page number, source file path) for traceability."
    )
    add_heading_styled(doc, "4.2 Text Splitting", 2)
    doc.add_paragraph(
        "Raw text is split using RecursiveCharacterTextSplitter with the following parameters:"
    )
    add_table(doc, ["Parameter", "Value", "Rationale"], [
        ("chunk_size", "1500", "Balances context richness with embedding quality"),
        ("chunk_overlap", "200", "Ensures continuity across chunk boundaries"),
        ("length_function", "len", "Python's built-in character count"),
    ])
    doc.add_paragraph(
        "The recursive splitter tries to split on paragraph breaks first, "
        "then sentences, then words — preserving semantic coherence."
    )
    add_heading_styled(doc, "4.3 Embedding & Storage", 2)
    doc.add_paragraph(
        "Each chunk is converted to a dense vector embedding using the Llama2 model's "
        "embedding layer via the Ollama API. These embeddings are stored in ChromaDB "
        "with the original text and metadata. The vector store is persisted to disk "
        "in the vectorDb/ directory, so documents survive across application restarts."
    )
    add_heading_styled(doc, "4.4 Deduplication", 2)
    doc.add_paragraph(
        "The application tracks which files have been ingested by inspecting the "
        "source metadata already present in ChromaDB. If a file has already been "
        "embedded, it is skipped to avoid duplicate entries."
    )

    # -- 5. RAG Query Pipeline --
    add_heading_styled(doc, "5. RAG Query Pipeline", 1)
    doc.add_paragraph(
        "When a user submits a question, the system executes the following pipeline:"
    )
    add_heading_styled(doc, "5.1 Retrieval", 2)
    doc.add_paragraph(
        "The user's question is embedded using the same Llama2 embedding model. "
        "ChromaDB performs a cosine similarity search and returns the top-4 most "
        "relevant document chunks. These chunks form the 'context' for the LLM."
    )
    add_heading_styled(doc, "5.2 Prompt Construction", 2)
    doc.add_paragraph("The prompt template used is:")
    doc.add_paragraph(
        'You are a knowledgeable chatbot, here to help with questions about '
        'the uploaded PDF documents.\n\n'
        'Context: {context}\n'
        'History: {history}\n\n'
        'User: {question}\n'
        'Chatbot:',
        style="No Spacing"
    )
    doc.add_paragraph(
        "The {context} placeholder is filled with retrieved document chunks, "
        "{history} with previous conversation turns, and {question} with the "
        "user's current query."
    )
    add_heading_styled(doc, "5.3 LLM Generation", 2)
    doc.add_paragraph(
        "The assembled prompt is sent to the Llama2 model running on Ollama. "
        "The response is streamed token-by-token to the Streamlit UI, providing "
        "a real-time typing effect. This improves perceived responsiveness, "
        "especially for longer answers."
    )

    # -- 6. Conversation Memory --
    add_heading_styled(doc, "6. Conversation Memory", 1)
    doc.add_paragraph(
        "The system uses LangChain's ConversationBufferMemory to maintain a "
        "full history of all user-assistant exchanges within a session. This enables "
        "follow-up questions, pronoun resolution ('What else did he do?'), and "
        "contextual continuity. The memory is stored in Streamlit's session state "
        "and resets when the browser tab is refreshed."
    )

    # -- 7. Project Structure --
    add_heading_styled(doc, "7. Project Structure", 1)
    structure = [
        ("app.py", "Main Streamlit application with RAG pipeline"),
        ("generate_docs.py", "This script — generates Word documentation"),
        ("requirements.txt", "Python package dependencies with version pins"),
        ("README.md", "Project overview and setup instructions"),
        (".gitignore", "Git ignore rules for user data and generated files"),
        ("LICENSE", "MIT License"),
        ("pdfFiles/", "Directory for uploaded PDF documents (git-ignored)"),
        ("vectorDb/", "ChromaDB persistent storage directory (git-ignored)"),
        ("docs/", "Generated documentation output"),
    ]
    add_table(doc, ["File / Directory", "Description"],
              [(f, d) for f, d in structure])

    # -- 8. Installation & Setup --
    add_heading_styled(doc, "8. Installation & Setup", 1)
    add_heading_styled(doc, "Prerequisites", 2)
    for item in ["Python 3.10+", "Conda package manager", "Ollama installed and running"]:
        doc.add_paragraph(item, style="List Bullet")

    add_heading_styled(doc, "Step-by-Step Setup", 2)
    setup_steps = [
        "Clone the repository:\n  git clone https://github.com/<username>/RAG-PDF-Chatbot.git",
        "Create and activate conda environment:\n  conda create -n rag_env python=3.10 -y\n  conda activate rag_env",
        "Install dependencies:\n  pip install -r requirements.txt",
        "Pull the Llama2 model:\n  ollama pull llama2",
        "Start the application:\n  streamlit run app.py",
    ]
    for s in setup_steps:
        doc.add_paragraph(s, style="List Number")

    # -- 9. Configuration --
    add_heading_styled(doc, "9. Configuration", 1)
    config = [
        ("base_url", "http://localhost:11434", "Ollama server address"),
        ("model", "llama2", "LLM model name (can be changed to mistral, llama3, etc.)"),
        ("chunk_size", "1500", "Characters per text chunk"),
        ("chunk_overlap", "200", "Overlap between consecutive chunks"),
        ("search_kwargs.k", "4", "Number of retrieved chunks per query"),
    ]
    add_table(doc, ["Parameter", "Default", "Description"],
              [(p, d, desc) for p, d, desc in config])

    # -- 10. Usage Guide --
    add_heading_styled(doc, "10. Usage Guide", 1)
    usage = [
        "Open the application in your browser at http://localhost:8501.",
        "Click 'Upload a PDF file' and select a PDF document.",
        "Wait for the processing indicator to complete.",
        "Type a question about the document in the chat input box.",
        "Read the AI-generated response grounded in your document.",
        "Ask follow-up questions — the chatbot remembers conversation context.",
    ]
    for u in usage:
        doc.add_paragraph(u, style="List Number")

    # -- 11. Troubleshooting --
    add_heading_styled(doc, "11. Troubleshooting", 1)
    issues = [
        ("Ollama connection refused", "Ensure Ollama is running: ollama serve"),
        ("Empty responses", "Check that PDFs contain extractable text (not scanned images)"),
        ("Slow responses", "Llama2 is resource-intensive; consider using a smaller model like tinyllama"),
        ("Import errors", "Ensure all packages are installed: pip install -r requirements.txt"),
        ("ChromaDB errors", "Delete vectorDb/ folder and re-upload PDFs to rebuild the index"),
    ]
    add_table(doc, ["Issue", "Solution"], [(i, s) for i, s in issues])

    # -- 12. Future Improvements --
    add_heading_styled(doc, "12. Future Improvements", 1)
    improvements = [
        "Support for multiple file formats (DOCX, TXT, HTML)",
        "Model selection dropdown in the UI",
        "Chunk visualization — show which document sections were used for each answer",
        "Authentication and multi-user support",
        "GPU acceleration for faster inference",
        "Export chat history as PDF or Markdown",
    ]
    for imp in improvements:
        doc.add_paragraph(imp, style="List Bullet")

    return doc


if __name__ == "__main__":
    os.makedirs("docs", exist_ok=True)
    doc = build_document()
    output_path = os.path.join("docs", "RAG_PDF_Chatbot_Documentation.docx")
    doc.save(output_path)
    print(f"Documentation generated: {output_path}")
